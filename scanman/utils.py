from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Cm
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import re
ip_regex = re.compile('^([0-9]+\.){3}[0-9]+.html$')
xlsx_regex = re.compile('^.*\.xlsx$')
true_ip_regex = re.compile('^([0-9]+\.){3}[0-9]+$')



def html_names_of_path(path):
  path = path + '/' if not path.endwith('/') else path
  filenames = os.listdir(path)
  filenames = list(filter(lambda x: ".html" in x, filenames))
  filenames = [path+_ for _ in filenames]
  return filenames

def recursive_html_names_of_path(path):
  ret = []
  for current_dir, sub_dirs, file_names in os.walk(path):
    for name in file_names:
      if ip_regex.match(name):
        ret.append(os.path.join(current_dir, name))
  return ret

def recursive_xlsx_names_of_path(path):
  ret = []
  for current_dir, sub_dirs, file_names in os.walk(path):
    for name in file_names:
      if xlsx_regex.match(name):
        ret.append(os.path.join(current_dir, name))
  return ret

 



def gadget_fill_cell(row, fields: list):
  for i in range(len(fields)):
    # set row height
    row.height = Cm(0.8)
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    # set vertical center
    row.cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # set horizontal center
    paragraph = row.cells[i].paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # write content to cell
    paragraph.add_run(str(fields[i]))


def doc_add_comment(doc, comment: str):
  doc.add_paragraph().text = comment



def concat_path(path_head, path_tail):
  path_head = path_head.strip('/')
  path_tail = path_tail.lstrip('/')
  return path_head + '/' + path_tail


def gadget_fill_cell_super(cells, fields):
  for i in range(len(fields)):
    cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    paragraph = cells[i].paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run(str(fields[i]))


def gadget_set_row_height(rows):
  for row in rows:
    row.height = Cm(0.8)
    row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST


def singleton(class_):
  instances = {}

  def getinstance(*args, **kwargs):
    if class_ not in instances:
      instances[class_] = class_(*args, **kwargs)
    return instances[class_]
  return getinstance


def dedup(records, func):
  res = []
  cache = []
  for record in records:
    if func(record) not in cache:
      res.append(record)
      cache.append(func(record))
  return res
